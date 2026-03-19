# Ferramenta: Extrair Texto de Word (.docx)
# Usa python-docx (leve, CPU)

import sys
import os
import json
from pathlib import Path

sys.path.append(str(Path(__file__).parent.parent / "00_CORE"))
from src.modulos.utils import InterfaceBase, Utils
from src.config.config import PASTA_SAIDAS

import docx
import customtkinter as ctk
from tkinter import filedialog, messagebox
import threading
import zipfile
from lxml import etree

class FerramentaWordparaTexto:
    def __init__(self):
        self.documento = None
        self.caminho_docx = None
        self.info = {}
    
    def abrir_word(self, caminho):
        """Abre um arquivo .docx"""
        try:
            self.documento = docx.Document(caminho)
            self.caminho_docx = caminho
            
            # informações bsicas
            self.info = {
                "paragrafos": len(self.documento.paragraphs),
                "tabelas": len(self.documento.tables),
                "secoes": len(self.documento.sections)
            }
            
            # Tenta extrair propriedades
            try:
                with zipfile.ZipFile(caminho, 'r') as docx_zip:
                    with docx_zip.open('docProps/core.xml') as f:
                        tree = etree.parse(f)
                        root = tree.getroot()
                        
                        for elem in root:
                            tag = elem.tag.split('}')[-1]
                            if tag == 'title':
                                self.info['titulo'] = elem.text
                            elif tag == 'creator':
                                self.info['autor'] = elem.text
                            elif tag == 'description':
                                self.info['descricao'] = elem.text
            except:
                pass
            
            return True, self.info
        except Exception as e:
            return False, str(e)
    
    def extrair_texto(self, incluir_tabelas=True):
        """Extrai todo o texto do documento"""
        if self.documento is None:
            return None, "Documento no carregado"
        
        try:
            texto_completo = []
            
            # Pargrafos
            for i, para in enumerate(self.documento.paragraphs):
                if para.text.strip():
                    texto_completo.append(para.text)
            
            # Tabelas
            if incluir_tabelas:
                for i, tabela in enumerate(self.documento.tables):
                    texto_completo.append(f"\n--- Tabela {i+1} ---")
                    for linha in tabela.rows:
                        linha_texto = [celula.text for celula in linha.cells]
                        texto_completo.append(" | ".join(linha_texto))
            
            return "\n".join(texto_completo), "Sucesso"
        except Exception as e:
            return None, str(e)
    
    def extrair_por_estilo(self, estilo):
        """Extrai textos com um estilo específico"""
        if self.documento is None:
            return None, "Documento no carregado"
        
        try:
            textos = []
            for para in self.documento.paragraphs:
                if para.style.name == estilo and para.text.strip():
                    textos.append(para.text)
            
            return textos, "Sucesso"
        except Exception as e:
            return None, str(e)
    
    def extrair_titulos(self):
        """Extrai textos que parecem ttulos"""
        if self.documento is None:
            return None, "Documento no carregado"
        
        try:
            titulos = []
            estilos_titulo = ['Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Ttulo', 'Subttulo']
            
            for para in self.documento.paragraphs:
                if para.style.name in estilos_titulo and para.text.strip():
                    titulos.append({
                        "estilo": para.style.name,
                        "texto": para.text
                    })
            
            return titulos, "Sucesso"
        except Exception as e:
            return None, str(e)
    
    def extrair_tabelas(self):
        """Extrai todas as tabelas em formato estruturado"""
        if self.documento is None:
            return None, "Documento no carregado"
        
        try:
            tabelas = []
            for i, tabela in enumerate(self.documento.tables):
                dados_tabela = []
                for linha in tabela.rows:
                    linha_dados = [celula.text for celula in linha.cells]
                    dados_tabela.append(linha_dados)
                
                tabelas.append({
                    "indice": i,
                    "linhas": len(tabela.rows),
                    "colunas": len(tabela.columns),
                    "dados": dados_tabela
                })
            
            return tabelas, "Sucesso"
        except Exception as e:
            return None, str(e)
    
    def extrair_comentarios(self):
        """Extrai comentários do documento Word via python-docx (namespace ooxml)."""
        if not self.doc:
            return [], "Documento não carregado"
        try:
            from docx.oxml.ns import qn
            comentarios = []
            # Comentários ficam em document.part.comments_part (se existir)
            try:
                comments_part = self.doc.part.comments_part
                if comments_part is None:
                    return [], "Documento sem comentários"
                for comment in comments_part._element.findall(qn("w:comment")):
                    autor = comment.get(qn("w:author"), "Desconhecido")
                    data  = comment.get(qn("w:date"), "")
                    texto_partes = [
                        r.text for r in comment.iter()
                        if r.tag == qn("w:t") and r.text
                    ]
                    texto = " ".join(texto_partes).strip()
                    comentarios.append({
                        "autor": autor,
                        "data": data,
                        "texto": texto,
                    })
            except AttributeError:
                # Fallback: varrer XML manualmente
                from lxml import etree
                ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                body = self.doc.element.body
                for el in body.iter():
                    if el.tag == f"{{{ns}}}commentReference":
                        comentarios.append({"ref": el.get(f"{{{ns}}}id", ""), "texto": "(referência no corpo)"})
            return comentarios, f"{len(comentarios)} comentário(s) extraído(s)"
        except Exception as e:
            return [], f"Erro ao extrair comentários: {e}"

class InterfaceWordparaTexto(InterfaceBase):
    def __init__(self):
        super().__init__(" Extrair Texto de Word", "800x700")
        self.ferramenta = FerramentaWordparaTexto()
        self.info_doc = None
        self.texto_extraido = None
        self.setup_interface()
    
    def setup_interface(self):
        titulo = ctk.CTkLabel(
            self.frame,
            text=" Extrair Texto de Arquivos Word (.docx)",
            font=("Arial", 24, "bold")
        )
        titulo.pack(pady=10)
        
        # Seleo
        self.frame_arquivo = ctk.CTkFrame(self.frame)
        self.frame_arquivo.pack(pady=10, padx=10, fill="x")
        
        self.btn_word = ctk.CTkButton(
            self.frame_arquivo,
            text=" Selecionar Word",
            command=self.selecionar_word,
            width=150,
            height=40
        )
        self.btn_word.pack(side="left", padx=5)
        
        self.lbl_arquivo = ctk.CTkLabel(
            self.frame_arquivo,
            text="Nenhum arquivo selecionado"
        )
        self.lbl_arquivo.pack(side="left", padx=10)
        
        # informações
        self.frame_info = ctk.CTkFrame(self.frame)
        self.frame_info.pack(pady=10, padx=10, fill="x")
        
        self.texto_info = ctk.CTkTextbox(self.frame_info, height=80)
        self.texto_info.pack(pady=5, padx=5, fill="x")
        
        # Abas
        self.tabview = ctk.CTkTabview(self.frame)
        self.tabview.pack(pady=10, padx=10, fill="both", expand=True)
        
        # Aba: Texto Completo
        self.tab_completo = self.tabview.add("Texto Completo")
        self.setup_tab_completo()
        
        # Aba: Ttulos
        self.tab_titulos = self.tabview.add("Ttulos")
        self.setup_tab_titulos()
        
        # Aba: Tabelas
        self.tab_tabelas = self.tabview.add("Tabelas")
        self.setup_tab_tabelas()
        
        # Aba: Por Estilo
        self.tab_estilo = self.tabview.add("Por Estilo")
        self.setup_tab_estilo()
        
        # Barra de progresso
        self.progress = ctk.CTkProgressBar(self.frame, width=400)
        self.progress.pack(pady=5)
        self.progress.set(0)
    
    def setup_tab_completo(self):
        self.frame_opcoes = ctk.CTkFrame(self.tab_completo)
        self.frame_opcoes.pack(pady=5, fill="x")
        
        self.incluir_tabelas = ctk.BooleanVar(value=True)
        self.chk_tabelas = ctk.CTkCheckBox(
            self.frame_opcoes,
            text="Incluir tabelas",
            variable=self.incluir_tabelas
        )
        self.chk_tabelas.pack(side="left", padx=5)
        
        self.btn_extrair = ctk.CTkButton(
            self.tab_completo,
            text=" Extrair Texto Completo",
            command=self.extrair_completo,
            width=150,
            height=35,
            fg_color="green",
            state="disabled"
        )
        self.btn_extrair.pack(pady=10)
        
        self.texto_resultado = ctk.CTkTextbox(self.tab_completo, height=200)
        self.texto_resultado.pack(pady=5, padx=5, fill="both", expand=True)
        
        self.frame_botoes = ctk.CTkFrame(self.tab_completo)
        self.frame_botoes.pack(pady=5)
        
        self.btn_copiar = ctk.CTkButton(
            self.frame_botoes,
            text=" Copiar",
            command=self.copiar_texto,
            width=100,
            state="disabled"
        )
        self.btn_copiar.pack(side="left", padx=5)
        
        self.btn_salvar = ctk.CTkButton(
            self.frame_botoes,
            text=" Salvar TXT",
            command=self.salvar_texto,
            width=100,
            state="disabled"
        )
        self.btn_salvar.pack(side="left", padx=5)
    
    def setup_tab_titulos(self):
        self.btn_titulos = ctk.CTkButton(
            self.tab_titulos,
            text=" Extrair Ttulos",
            command=self.extrair_titulos,
            width=150,
            height=35,
            fg_color="blue",
            state="disabled"
        )
        self.btn_titulos.pack(pady=10)
        
        self.texto_titulos = ctk.CTkTextbox(self.tab_titulos, height=250)
        self.texto_titulos.pack(pady=5, padx=5, fill="both", expand=True)
    
    def setup_tab_tabelas(self):
        self.btn_tabelas = ctk.CTkButton(
            self.tab_tabelas,
            text=" Extrair Tabelas",
            command=self.extrair_tabelas,
            width=150,
            height=35,
            fg_color="orange",
            state="disabled"
        )
        self.btn_tabelas.pack(pady=10)
        
        self.texto_tabelas = ctk.CTkTextbox(self.tab_tabelas, height=250)
        self.texto_tabelas.pack(pady=5, padx=5, fill="both", expand=True)
    
    def setup_tab_estilo(self):
        self.frame_estilo = ctk.CTkFrame(self.tab_estilo)
        self.frame_estilo.pack(pady=10, padx=5, fill="x")
        
        self.lbl_estilo = ctk.CTkLabel(self.frame_estilo, text="Estilo:")
        self.lbl_estilo.pack(side="left", padx=5)
        
        self.entry_estilo = ctk.CTkEntry(
            self.frame_estilo,
            placeholder_text="Ex: Heading 1, Title, Normal",
            width=200
        )
        self.entry_estilo.pack(side="left", padx=5)
        
        self.btn_estilo = ctk.CTkButton(
            self.tab_estilo,
            text=" Extrair por Estilo",
            command=self.extrair_por_estilo,
            width=150,
            height=35,
            fg_color="purple",
            state="disabled"
        )
        self.btn_estilo.pack(pady=10)
        
        self.texto_estilo = ctk.CTkTextbox(self.tab_estilo, height=250)
        self.texto_estilo.pack(pady=5, padx=5, fill="both", expand=True)
    
    def selecionar_word(self):
        caminho = filedialog.askopenfilename(
            title="Selecione um arquivo Word",
            filetypes=[("Word", "*.docx")]
        )
        if caminho:
            self.lbl_arquivo.configure(text=f"Arquivo: {Path(caminho).name}")
            
            sucesso, info = self.ferramenta.abrir_word(caminho)
            if sucesso:
                self.info_doc = info
                
                info_texto = f"Pargrafos: {info.get('paragrafos', 'N/A')}\n"
                info_texto += f"Tabelas: {info.get('tabelas', 'N/A')}\n"
                info_texto += f"Ttulo: {info.get('titulo', 'N/A')}\n"
                info_texto += f"Autor: {info.get('autor', 'N/A')}"
                
                self.texto_info.delete('1.0', 'end')
                self.texto_info.insert('1.0', info_texto)
                
                # Ativa botes
                self.btn_extrair.configure(state="normal")
                self.btn_titulos.configure(state="normal")
                self.btn_tabelas.configure(state="normal")
                self.btn_estilo.configure(state="normal")
            else:
                self.utils.mostrar_erro("Erro", info)
    
    def extrair_completo(self):
        def extrair_thread():
            self.btn_extrair.configure(state="disabled", text=" Extraindo...")
            self.progress.set(0.3)
            
            texto, msg = self.ferramenta.extrair_texto(
                incluir_tabelas=self.incluir_tabelas.get()
            )
            
            self.progress.set(0.8)
            
            if texto:
                self.texto_extraido = texto
                self.texto_resultado.delete('1.0', 'end')
                self.texto_resultado.insert('1.0', texto[:1000] + "...\n\n(Preview)")
                self.btn_copiar.configure(state="normal")
                self.btn_salvar.configure(state="normal")
                self.utils.mostrar_info("Sucesso", "Texto extrado!")
            else:
                self.utils.mostrar_erro("Erro", msg)
            
            self.progress.set(1)
            self.btn_extrair.configure(state="normal", text=" Extrair Texto Completo")
        
        threading.Thread(target=extrair_thread).start()
    
    def extrair_titulos(self):
        def titulos_thread():
            self.btn_titulos.configure(state="disabled", text=" Extraindo...")
            
            titulos, msg = self.ferramenta.extrair_titulos()
            
            if titulos:
                self.texto_titulos.delete('1.0', 'end')
                for titulo in titulos:
                    self.texto_titulos.insert('end', f"[{titulo['estilo']}] {titulo['texto']}\n")
            else:
                self.texto_titulos.insert('end', "Nenhum ttulo encontrado")
            
            self.btn_titulos.configure(state="normal", text=" Extrair Ttulos")
        
        threading.Thread(target=titulos_thread).start()
    
    def extrair_tabelas(self):
        def tabelas_thread():
            self.btn_tabelas.configure(state="disabled", text=" Extraindo...")
            
            tabelas, msg = self.ferramenta.extrair_tabelas()
            
            if tabelas:
                self.texto_tabelas.delete('1.0', 'end')
                for tabela in tabelas:
                    self.texto_tabelas.insert('end', f"--- Tabela {tabela['indice']+1} ({tabela['linhas']}x{tabela['colunas']}) ---\n")
                    for linha in tabela['dados']:
                        self.texto_tabelas.insert('end', " | ".join(linha) + "\n")
                    self.texto_tabelas.insert('end', "\n")
            else:
                self.texto_tabelas.insert('end', "Nenhuma tabela encontrada")
            
            self.btn_tabelas.configure(state="normal", text=" Extrair Tabelas")
        
        threading.Thread(target=tabelas_thread).start()
    
    def extrair_por_estilo(self):
        def estilo_thread():
            self.btn_estilo.configure(state="disabled", text=" Extraindo...")
            
            estilo = self.entry_estilo.get().strip()
            if not estilo:
                self.utils.mostrar_erro("Erro", "Digite um estilo")
                self.btn_estilo.configure(state="normal", text=" Extrair por Estilo")
                return
            
            textos, msg = self.ferramenta.extrair_por_estilo(estilo)
            
            self.texto_estilo.delete('1.0', 'end')
            if textos:
                for i, texto in enumerate(textos, 1):
                    self.texto_estilo.insert('end', f"{i}. {texto}\n")
            else:
                self.texto_estilo.insert('end', f"Nenhum texto encontrado com estilo '{estilo}'")
            
            self.btn_estilo.configure(state="normal", text=" Extrair por Estilo")
        
        threading.Thread(target=estilo_thread).start()
    
    def copiar_texto(self):
        if self.texto_extraido:
            self.frame.clipboard_clear()
            self.frame.clipboard_append(self.texto_extraido[:10000])
            self.utils.mostrar_info("Copiado", "Texto copiado")
    
    def salvar_texto(self):
        if self.texto_extraido:
            caminho = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Texto", "*.txt")]
            )
            if caminho:
                with open(caminho, 'w', encoding='utf-8') as f:
                    f.write(self.texto_extraido)
                self.utils.mostrar_info("Sucesso", f"Texto salvo em:\n{caminho}")

if __name__ == "__main__":
    app = InterfaceWordparaTexto()
    app.rodar()
